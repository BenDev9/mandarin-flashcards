from docx import Document
from formating import format_card, format_table

TABLE_ROWS = 5
TEMPLATE_TEXT = "\n\n\n"

BLANK_CARD = {"Front": TEMPLATE_TEXT * 2 + "\n", "Back": TEMPLATE_TEXT * 2 + "\n"}


def generate(pairs: list):
    out = Document()
    fronts = out.add_table(rows=1, cols=2)
    backs = out.add_table(rows=1, cols=2)

    frontCells, backCells = fronts.rows[0].cells, backs.rows[0].cells

    pair = pairs[0]

    frontCells[0].text, backCells[1].text = format_card(pair[0])
    frontCells[1].text, backCells[0].text = format_card(pair[1])

    for i in range(1, len(pairs)):
        if i % TABLE_ROWS == 0:
            # Set table styles
            format_table(fronts, backs)

            fronts = out.add_table(rows=1, cols=2)
            backs = out.add_table(rows=1, cols=2)

            frontCells, backCells = fronts.rows[0].cells, backs.rows[0].cells

            i += 1
            if i == len(pairs):
                continue
            pair = pairs[i]

            frontCells[0].text, backCells[0].text = format_card(pair[0])
            frontCells[1].text, backCells[1].text = format_card(pair[1])
        else:
            frontCells, backCells = fronts.add_row().cells, backs.add_row().cells

            pair = pairs[i]

            frontCells[0].text, backCells[0].text = format_card(pair[0])

            if len(pair) > 1:
                frontCells[1].text, backCells[1].text = format_card(pair[1])
            else:
                frontCells[1].text, backCells[1].text = format_card(BLANK_CARD)

    # Set table styles
    format_table(fronts, backs)

    return out
